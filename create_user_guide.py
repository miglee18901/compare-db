from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from pathlib import Path

OUT = Path('Tai_lieu_huong_dan_su_dung_CompareDB.docx')

BLUE = '2E74B5'
DARK = '1F4D78'
NAVY = '0B2545'
LIGHT = 'E8EEF5'
GRAY = 'F4F6F9'

def set_run_font(run, name='Calibri', size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')

def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(dxa))
    tc_w.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def write_cell(cell, text, bold=False, color=None, size=10.5):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_geometry(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, BLUE)
        write_cell(cell, header, bold=True, color='FFFFFF', size=10)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            write_cell(cell, value, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.38 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r)
    return p

def add_step(doc, title, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(title + ': ')
    set_run_font(r, bold=True, color=DARK)
    r = p.add_run(text)
    set_run_font(r)

def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + ' ')
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)
    return p

def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [('Heading 1',16,BLUE,16,8), ('Heading 2',13,BLUE,11,6), ('Heading 3',12,DARK,8,4)]:
        st = styles[name]
        st.font.name = 'Calibri'
        st._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    # Header/footer
    h = section.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = h.add_run('HƯỚNG DẪN SỬ DỤNG - ĐỐI CHIẾU DỮ LIỆU')
    set_run_font(hr, size=8.5, color='666666')
    f = section.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run('Tài liệu hướng dẫn người sử dụng')
    set_run_font(fr, size=8.5, color='666666')

    # Opening block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('HƯỚNG DẪN SỬ DỤNG')
    set_run_font(r, size=12, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('Đối chiếu dữ liệu giữa hai môi trường')
    set_run_font(r, size=25, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run('Tài liệu dành cho người vận hành')
    set_run_font(r, size=13, color='555555', italic=True)

    add_note(doc, 'Mục đích:', 'Hỗ trợ kiểm tra dữ liệu giữa môi trường CRBT16M và CRBT21M, nhanh chóng nhận biết bản ghi thiếu hoặc có thông tin khác nhau.')

    doc.add_heading('1. Tổng quan', level=1)
    add_body(doc, 'Công cụ đọc danh sách các bảng cần kiểm tra, kết nối đến hai môi trường dữ liệu và lập báo cáo kết quả. Người sử dụng chọn một trong hai cách kiểm tra: đối chiếu chi tiết từng bản ghi hoặc chỉ đếm tổng số bản ghi của từng bảng.')
    add_table(doc, ['Nội dung', 'Ý nghĩa'], [
        ('Môi trường 1', 'CRBT16M - nguồn dữ liệu thứ nhất.'),
        ('Môi trường 2', 'CRBT21M - nguồn dữ liệu thứ hai.'),
        ('Đối chiếu chi tiết', 'Kiểm tra từng bản ghi theo cột nhận diện và so từng thông tin cần kiểm tra.'),
        ('Đếm số bản ghi', 'So sánh tổng số bản ghi của mỗi bảng, không đi sâu vào nội dung từng bản ghi.')
    ], [2800, 6560])

    doc.add_heading('2. Mô tả cấu hình', level=1)
    add_body(doc, 'Trước khi chạy, người vận hành chuẩn bị bốn nhóm thông tin: nơi kết nối của hai môi trường, chế độ kiểm tra, nơi lưu báo cáo và danh sách bảng cần đối chiếu.')
    doc.add_heading('2.1. Chọn chế độ kiểm tra và nơi lưu kết quả', level=2)
    add_body(doc, 'Các lựa chọn này nằm trong tệp config.properties.')
    add_table(doc, ['Tên lựa chọn', 'Cách sử dụng', 'Ví dụ'], [
        ('MODE', 'Chọn cách kiểm tra. Giá trị 1 là đối chiếu chi tiết; giá trị 2 là đếm số bản ghi.', 'MODE=1'),
        ('PATH_STATISTICS_FILE', 'Chỉ định thư mục lưu báo cáo. Dấu chấm (.) nghĩa là lưu tại thư mục đang chạy công cụ.', 'PATH_STATISTICS_FILE=.'),
        ('BATCH_SIZE', 'Số lượng bản ghi được đọc theo từng lượt. Thường giữ giá trị 1000.', 'BATCH_SIZE=1000')
    ], [2600, 4760, 2000])
    add_note(doc, 'Lưu ý:', 'Chỉ dùng MODE=1 khi cần biết chính xác bản ghi và trường thông tin nào khác nhau. Dùng MODE=2 khi chỉ cần kiểm tra nhanh số lượng.')

    doc.add_heading('2.2. Danh sách bảng cần kiểm tra', level=2)
    add_body(doc, 'Danh sách nằm trong tệp tableList.txt. Mỗi dòng tương ứng với một bảng, theo mẫu: TÊN_BẢNG|CỘT_NHẬN_DIỆN|CỘT_BỎ_QUA.')
    add_table(doc, ['Ví dụ cấu hình', 'Diễn giải'], [
        ('SUBS_INFO|MSISDN|ignoreColumn1,ignoreColumn2', 'Kiểm tra bảng SUBS_INFO; dùng MSISDN để nhận diện từng bản ghi; không xét hai cột ignoreColumn1 và ignoreColumn2.'),
        ('PROD_SPEC|PROD_SPEC_ID|', 'Kiểm tra bảng PROD_SPEC; dùng PROD_SPEC_ID để nhận diện từng bản ghi; không có cột nào được bỏ qua.')
    ], [4300, 5060])
    add_bullet(doc, 'Cột nhận diện phải giúp xác định duy nhất một bản ghi trong bảng.')
    add_bullet(doc, 'Các cột bỏ qua được dùng khi chấp nhận việc hai môi trường có thể khác nhau ở các thông tin này.')
    add_bullet(doc, 'Có thể thêm nhiều bảng bằng cách thêm mỗi bảng trên một dòng mới.')
    add_bullet(doc, 'Dòng trống hoặc dòng bắt đầu bằng dấu # sẽ không được kiểm tra.')

    doc.add_heading('2.3. Thông tin kết nối hai môi trường', level=2)
    add_body(doc, 'Hai tệp trong thư mục crbt16m và crbt21m chứa địa chỉ, tên môi trường, tài khoản và mật khẩu để truy cập dữ liệu. Người vận hành cần bảo đảm các thông tin này đúng và tài khoản được cấp quyền xem dữ liệu của các bảng cần kiểm tra.')
    add_note(doc, 'Bảo mật:', 'Không chia sẻ tệp chứa tài khoản và mật khẩu cho người không có thẩm quyền. Thay đổi mật khẩu trong hai tệp kết nối khi mật khẩu môi trường thay đổi.')

    doc.add_heading('3. Luồng đối chiếu chi tiết từng bản ghi/cột', level=1)
    add_body(doc, 'Chế độ này được chọn khi MODE=1. Mục tiêu là xác định rõ bản ghi nào khác nhau, thiếu ở môi trường nào và thông tin nào không giống nhau.')
    doc.add_heading('3.1. Các bước thực hiện', level=2)
    add_step(doc, 'Đọc danh sách bảng', 'Công cụ lần lượt lấy từng dòng hợp lệ trong tableList.txt.')
    add_step(doc, 'Kiểm tra điều kiện', 'Kiểm tra bảng có tồn tại ở cả hai môi trường; cột nhận diện và các cột bỏ qua có đúng hay không.')
    add_step(doc, 'Xác định thông tin cần kiểm tra', 'Bỏ cột nhận diện và các cột được khai báo bỏ qua. Chỉ kiểm tra các thông tin có mặt ở cả hai môi trường.')
    add_step(doc, 'Đọc dữ liệu theo từng lượt', 'Dữ liệu được đọc theo số lượng đã chọn ở BATCH_SIZE để phù hợp với bảng có nhiều bản ghi.')
    add_step(doc, 'Ghép bản ghi theo cột nhận diện', 'Mỗi bản ghi ở CRBT16M được đối chiếu với bản ghi có cùng giá trị nhận diện ở CRBT21M.')
    add_step(doc, 'So sánh thông tin', 'Nếu hai bản ghi cùng tồn tại, công cụ so từng thông tin cần kiểm tra và ghi lại các điểm khác nhau.')
    add_step(doc, 'Lập báo cáo', 'Báo cáo hiển thị tổng số bản ghi đã kiểm tra, số khớp, số không khớp và chi tiết chênh lệch.')

    doc.add_heading('3.2. Các tình huống trong kết quả', level=2)
    add_table(doc, ['Tình huống', 'Kết quả ghi nhận'], [
        ('Bản ghi có ở cả hai nơi và mọi thông tin cần kiểm tra giống nhau', 'Được tính là khớp.'),
        ('Bản ghi có ở cả hai nơi nhưng một hoặc nhiều thông tin khác nhau', 'Được tính là không khớp; báo cáo nêu từng thông tin khác nhau.'),
        ('Bản ghi chỉ có ở CRBT16M', 'Được tính là không khớp; phía CRBT21M được ghi là không tồn tại.'),
        ('Bản ghi chỉ có ở CRBT21M', 'Được tính là không khớp; phía CRBT16M được ghi là không tồn tại.'),
        ('Bảng hoặc cột cấu hình không đúng', 'Báo lỗi cấu hình cho bảng đó và bỏ qua việc đối chiếu bảng này.')
    ], [3450, 5910])
    doc.add_heading('3.3. Cách đọc báo cáo chi tiết', level=2)
    add_body(doc, 'Mỗi bảng có một dòng tổng hợp. Ví dụ:')
    add_note(doc, 'Ví dụ:', 'TABLE = SUBS_INFO, TOTAL = 4, MATCH = 1, NOT_MATCH = 3')
    add_bullet(doc, 'TOTAL: tổng số bản ghi nhận diện được khi gộp dữ liệu của hai môi trường.')
    add_bullet(doc, 'MATCH: số bản ghi có ở cả hai môi trường và các thông tin cần kiểm tra đều giống nhau.')
    add_bullet(doc, 'NOT_MATCH: số bản ghi thiếu ở một bên hoặc có ít nhất một thông tin khác nhau.')
    add_body(doc, 'Phần COMPARE bên dưới dòng tổng hợp cho biết chi tiết. Ví dụ:')
    add_note(doc, 'Ví dụ:', 'MSISDN key = 2\nSINGER (CRBT16M) = S2_old, SINGER (CRBT21M) = S2_new')
    add_body(doc, 'Ví dụ trên cho biết bản ghi có MSISDN bằng 2 tồn tại ở cả hai môi trường, nhưng thông tin SINGER khác nhau.')

    doc.add_heading('4. Luồng đếm số bản ghi mỗi bảng', level=1)
    add_body(doc, 'Chế độ này được chọn khi MODE=2. Đây là cách kiểm tra nhanh để biết số lượng bản ghi của cùng một bảng có bằng nhau hay không. Chế độ này không cho biết bản ghi cụ thể nào bị thiếu hoặc khác nội dung.')
    doc.add_heading('4.1. Các bước thực hiện', level=2)
    add_step(doc, 'Đọc danh sách bảng', 'Công cụ lấy từng bảng trong tableList.txt.')
    add_step(doc, 'Kiểm tra bảng', 'Kiểm tra bảng có tồn tại ở CRBT16M và CRBT21M hay không.')
    add_step(doc, 'Đếm số lượng', 'Đếm tổng số bản ghi của bảng tại từng môi trường.')
    add_step(doc, 'Ghi kết quả', 'Xuất một dòng cho mỗi bảng, nêu số lượng tại CRBT16M và CRBT21M.')
    doc.add_heading('4.2. Cách đọc báo cáo đếm số lượng', level=2)
    add_note(doc, 'Ví dụ:', 'TABLE = SUBS_INFO, CRBT16M = 100, CRBT21M = 105')
    add_bullet(doc, 'Hai số bằng nhau: số lượng bản ghi của bảng là như nhau.')
    add_bullet(doc, 'Hai số khác nhau: cần chuyển sang MODE=1 nếu cần tìm chính xác bản ghi bị thiếu hoặc có khác biệt.')
    add_bullet(doc, 'Nếu bảng không tồn tại ở một môi trường, báo cáo sẽ ghi rõ môi trường và tên bảng gặp vấn đề.')

    doc.add_heading('5. Sau khi chạy', level=1)
    add_body(doc, 'Kết quả được tạo dưới dạng tệp có tên result_YYYYMMDDHHMMSS.txt tại thư mục đã chọn ở PATH_STATISTICS_FILE. Đường dẫn đầy đủ của tệp cũng được hiển thị khi công cụ chạy xong.')
    add_table(doc, ['Khi cần', 'Nên chọn'], [
        ('Kiểm tra nhanh tổng số dữ liệu giữa hai môi trường', 'MODE=2'),
        ('Tìm bản ghi thiếu hoặc thông tin khác nhau', 'MODE=1'),
        ('Không muốn xét một số thông tin thường xuyên thay đổi', 'Khai báo các thông tin đó ở phần cột bỏ qua trong tableList.txt')
    ], [4900, 4460])
    add_note(doc, 'Khuyến nghị:', 'Bắt đầu bằng MODE=2 để phát hiện chênh lệch số lượng. Khi cần xác minh nguyên nhân, chuyển sang MODE=1 cho các bảng cần điều tra.')

    doc.core_properties.title = 'Hướng dẫn sử dụng công cụ đối chiếu dữ liệu'
    doc.core_properties.subject = 'Cấu hình và luồng đối chiếu dữ liệu'
    doc.core_properties.author = 'CompareDB'
    doc.save(OUT)
    print(OUT.resolve())

if __name__ == '__main__':
    main()
